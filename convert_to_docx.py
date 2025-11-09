"""
Convert readme.md to a 1960s-style computer manual DOCX
Uses Courier font and vintage formatting
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_vintage_manual():
    """Create a DOCX with 1960s computer manual styling"""
    
    # Create document
    doc = Document()
    
    # Set narrow margins for vintage feel
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Read the markdown file
    with open('readme.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Process line by line
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines at start of document
        if not line.strip() and len(doc.paragraphs) <= 1:
            i += 1
            continue
        
        # Main title (# BBC BASEBALL...)
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].upper())
            run.font.name = 'Courier New'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.space_after = Pt(6)
        
        # Section headers (## INSTRUCTION MANUAL, etc.)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[3:].upper())
            run.font.name = 'Courier New'
            run.font.size = Pt(14)
            run.font.bold = True
            p.space_before = Pt(12)
            p.space_after = Pt(6)
        
        # Subsection headers (### CONGRATULATIONS!, etc.)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].upper())
            run.font.name = 'Courier New'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.underline = True
            p.space_before = Pt(12)
            p.space_after = Pt(6)
        
        # Horizontal rules
        elif line.strip().startswith('---'):
            p = doc.add_paragraph()
            run = p.add_run('=' * 60)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(12)
            p.space_after = Pt(12)
        
        # Code blocks
        elif line.strip().startswith('```'):
            # Start of code block
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # Add code block with box border effect using ASCII characters
            if code_lines:
                # Calculate max width for proper box formatting
                max_width = max(len(line) for line in code_lines) if code_lines else 60
                max_width = min(max_width, 68)  # Cap width to prevent wrapping
                
                # Top border
                p = doc.add_paragraph()
                border_width = max_width + 2
                run = p.add_run('+' + '-' * border_width + '+')
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.keep_together = True
                
                # Code content
                for code_line in code_lines:
                    p = doc.add_paragraph()
                    # Truncate or pad line to max width
                    if len(code_line) > max_width:
                        display_line = code_line[:max_width]
                    else:
                        display_line = code_line.ljust(max_width)
                    run = p.add_run('| ' + display_line + ' |')
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.keep_together = True
                
                # Bottom border
                p = doc.add_paragraph()
                run = p.add_run('+' + '-' * border_width + '+')
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.keep_together = True
                p.space_after = Pt(12)
        
        # Bold text in lists (**PLAYER ROSTER**, etc.)
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip()[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text.upper())
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.bold = True
            p.space_before = Pt(8)
            p.space_after = Pt(4)
        
        # Bullet points (starting with -)
        elif line.strip().startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            # Remove markdown formatting
            text = line.strip()[2:]
            # Handle bold within text
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            run = p.add_run(text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            p.paragraph_format.left_indent = Inches(0.5)
        
        # Regular paragraphs
        elif line.strip():
            p = doc.add_paragraph()
            
            # Process inline formatting
            text = line.strip()
            # Remove markdown links but keep text
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            
            # Split by bold markers
            parts = re.split(r'(\*\*.*?\*\*)', text)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.bold = True
                elif part:
                    # Handle inline code
                    code_parts = re.split(r'(`.*?`)', part)
                    for code_part in code_parts:
                        if code_part.startswith('`') and code_part.endswith('`'):
                            run = p.add_run(code_part[1:-1])
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
                        elif code_part:
                            run = p.add_run(code_part)
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
            
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Empty lines
        else:
            doc.add_paragraph()
        
        i += 1
    
    # Save the document
    doc.save('BBC_Baseball_Manual_1960s.docx')
    print("✓ Document created: BBC_Baseball_Manual_1960s.docx")
    print("  Font: Courier New (vintage typewriter style)")
    print("  Format: 1960s computer manual aesthetic")

if __name__ == '__main__':
    create_vintage_manual()
